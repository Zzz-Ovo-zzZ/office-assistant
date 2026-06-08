"""八字合婚报告 → Word v3.0 — 对标MD排版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                      'output', '八字合婚分析报告_20260608.docx')

doc = Document()

# ─── 页面 ───
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

GOLD = 'B8860B'; DARK = '8B6914'; BG = 'FFF9F0'
WH = 'FFFFFF'; GR = '888888'; BK = '333333'

# ─── helpers ───
def ea(): return qn('w:eastAsia')

def fr(run, sz=10.5, b=False, c=None, fn='微软雅黑'):
    run.font.name = fn; run.font.size = Pt(sz); run.bold = b
    run._element.rPr.rFonts.set(ea(), fn)
    if c: run.font.color.rgb = RGBColor.from_string(c)

def ap(text='', sz=10.5, b=False, c=None, al=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = al
    r = p.add_run(text); fr(r, sz, b, c); return p

def acres(text='', sz=10.5, b=False, c=None):
    return ap(text, sz, b, c, WD_ALIGN_PARAGRAPH.CENTER)

def tcell(cell, text, sz=9.5, b=False, c=None, al=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]; p.alignment = al
    pf = p.paragraph_format; pf.space_before = Pt(4); pf.space_after = Pt(4)
    r = p.add_run(str(text)); fr(r, sz, b, c); return r

def tcl(cell, text, sz=9.5, b=False, c=None):
    """left-aligned cell"""
    return tcell(cell, text, sz, b, c, WD_ALIGN_PARAGRAPH.LEFT)

def bgx(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def new_table(headers, rows, col_w=None, hdr_bg=GOLD, hdr_c=WH, zebra=BG):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; bgx(c, hdr_bg); tcell(c, h, sz=10, b=True, c=hdr_c)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            if ri % 2 == 0: bgx(c, zebra)
            is_left = (ci == len(row)-1 and len(row) > 1)
            if is_left: tcl(c, val, sz=9.5)
            else: tcell(c, val, sz=9.5)
    if col_w:
        for i, w in enumerate(col_w):
            for r in t.rows: r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def info_card(rows, col_w=[3,13]):
    """双列表格：标签 | 内容"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (k, v) in enumerate(rows):
        c0, c1 = t.rows[ri].cells[0], t.rows[ri].cells[1]
        if ri % 2 == 0: bgx(c0, BG); bgx(c1, BG)
        tcell(c0, k, sz=10, b=True, c=DARK)
        tcl(c1, v, sz=10)
    t.columns[0].width = Cm(col_w[0]); t.columns[1].width = Cm(col_w[1])
    doc.add_paragraph()
    return t

def sec(number, text, level=2):
    """段落标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(18); pf.space_after = Pt(10)
    r = p.add_run(text if not number else f'{number}、{text}')
    if level == 1:
        fr(r, sz=26, b=True, c=GOLD); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        fr(r, sz=15, b=True, c=GOLD)
    return p

def sub_sec(text):
    """小标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(14); pf.space_after = Pt(6)
    r = p.add_run(text); fr(r, sz=12, b=True, c=DARK)
    return p

def code_block(lines, bg=BG, sz=9.5):
    """灰底代码块 — 1列表格"""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; bgx(c, bg)
    for i, line in enumerate(lines):
        pa = c.paragraphs[0] if i == 0 else c.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = pa.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(2)
        r = pa.add_run(line); fr(r, sz=sz, fn='SimSun' if '─' not in line and '┌' not in line else 'SimSun')
    doc.add_paragraph()
    return t

def gold_box(lines, title_idx=0, sz=12, border_c=GOLD):
    """金边框提示框"""
    t = doc.add_table(rows=len(lines), cols=1)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, line in enumerate(lines):
        c = t.rows[i].cells[0]; bgx(c, BG)
        pa = c.paragraphs[0]; pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = pa.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(6)
        r = pa.add_run(line)
        is_title = (i == title_idx)
        fr(r, sz=sz+2 if is_title else sz, b=is_title, c=DARK if is_title else '5a3a10')
    doc.add_paragraph()
    return t

def pillar_table(title, gan, zhi, animals, traits):
    """八字排盘专用表"""
    doc.add_paragraph()
    acres(title, sz=13, b=True, c=DARK)
    t = doc.add_table(rows=5, cols=4)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = ['', '年柱', '月柱', '日柱', '时柱']
    for i, lab in enumerate(labels):
        tcell(t.rows[i].cells[0], lab, sz=10.5, b=True, c=DARK)

    heavenly = ['天干'] + gan
    earthly = ['地支'] + zhi
    animal = ['生肖'] + animals

    for i, label in enumerate(heavenly):
        c = t.rows[i].cells[1]; bgx(c, BG)
        tcell(c, label, sz=11, b=(i>0))
    for i, label in enumerate(earthly):
        c = t.rows[i].cells[2]; bgx(c, BG)
        tcell(c, label, sz=11, b=(i>0))
    for i, label in enumerate(animal):
        c = t.rows[i].cells[3]; bgx(c, BG)
        tcell(c, label, sz=10, c=DARK if i>0 else None)

    for col in [0,1,2,3]:
        for row in t.rows: row.cells[col].width = Cm([2, 3, 3, 2.5][col])

    # info rows
    info_rows = [
        ['日主', traits[0]],
        ['纳音', traits[1]],
        ['五行', traits[2]],
        ['特质', traits[3]],
    ]
    info_card(info_rows)
    return t

def merge_row(t, row_idx):
    """合并整行"""
    c = t.rows[row_idx].cells[0]
    for i in range(1, len(t.rows[row_idx].cells)):
        c.merge(t.rows[row_idx].cells[i])
    return c

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════
#  封面
# ═══════════════════════════════════════════════════════

doc.add_paragraph()
ap(); ap()
acres('八字合婚分析报告', sz=26, b=True, c=GOLD)
ap()
acres('男方  丙子年七月十四辰时  ×  女方  丁丑年正月十八申时', sz=11, c=GR)
acres('1996.08.27 08:00 福建  ──合──  1997.02.24 16:00 江西', sz=10.5)
ap(); ap()

gold_box(['综合评分', '⭐ 5.0 / 5.0  🟡🟡🟡🟡🟡', '上 等 婚 配'], sz=14)

page_break()

# ═══════════════════════════════════════════════════════
#  一、八字排盘
# ═══════════════════════════════════════════════════════

sec('一', '八字排盘')

# ── 男方 ──
pillar_table(
    '🚹 乾造 · 男方',
    ['丙', '丙', '丁', '甲'],
    ['子', '申', '酉', '辰'],
    ['鼠', '猴', '鸡', '龙'],
    ['丁火（灯烛之火 🔥）',
     '涧下水 · 山下火 · 山下火 · 覆灯火',
     '火 40% · 金 20% · 水 20% · 木 10% · 土 10%',
     '温和细腻，情感丰富，有文艺直觉。不是烈日（丙火），是冬夜炉火——不灼人，但恒久暖人']
)

# ── 女方 ──
pillar_table(
    '🚺 坤造 · 女方',
    ['丁', '壬', '戊', '庚'],
    ['丑', '寅', '戌', '申'],
    ['牛', '虎', '狗', '猴'],
    ['戊土（城墙之土 ⛰️）',
     '涧下水 · 金箔金 · 平地木 · 石榴木',
     '土 35% · 金 28% · 水 15% · 火 12% · 木 10%',
     '大气稳重，原则清晰，性格厚实。是城墙厚土——不随风飘，不轻易动摇']
)

# ═══════════════════════════════════════════════════════
#  二、合盘总评
# ═══════════════════════════════════════════════════════

sec('二', '合盘总评')
gold_box(['综合评分', '⭐ 5.0 / 5.0  🟡🟡🟡🟡🟡', '上 等 婚 配'], sz=15)

# ═══════════════════════════════════════════════════════
#  三、合盘五维详析
# ═══════════════════════════════════════════════════════

sec('三', '合盘五维详析')

# 维度一
sub_sec('🟡 维度一：生肖 · 子丑六合 —— ⭐⭐⭐⭐⭐ 满分')
info_card([
    ['关系', '子(水) + 丑(土) → 六合，合化土'],
    ['解读', '十二生肖中最和谐的组合之一。子丑不冲不害不刑，是天然互吸的磁场。民间讲"鼠牛配，一世情"，不是空话。'],
    ['现实', '你们之间有一种无需多言的默契。像两个拼图碎片，恰好是对方的缺口形状。'],
])

# 地支关系 — 用表格而非ASCII
t = doc.add_table(rows=5, cols=3); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (a, rel, b) in enumerate([
    ('子', '──六合──', '丑 ✅ 你俩'),
    ('子', '──三合──', '申辰（自力更生格局）'),
    ('子', '──冲───', '午 ❌ 不犯'),
    ('子', '──害───', '未 ❌ 不犯'),
    ('子', '──刑───', '卯 ❌ 不犯'),
]):
    for ci, v in enumerate([a, rel, b]):
        c = t.rows[ri].cells[ci]
        if ri % 2 == 0: bgx(c, BG)
        tcell(c, v, sz=9.5, b=(ci==1))
for r in t.rows:
    r.cells[0].width = Cm(3); r.cells[1].width = Cm(5); r.cells[2].width = Cm(8)
doc.add_paragraph()

# 维度二
sub_sec('🟡 维度二：日主 · 丁火生戊土 —— ⭐⭐⭐⭐⭐ 极佳')
info_card([
    ['关系', '男 丁火 ——相生——→ 女 戊土'],
    ['解读', '这是最理想的日主配法之一。男方天然地滋养女方——火燃烧自己，化为土。不是消耗型的"克"，而是创造型的"生"。'],
    ['现实', '他天生愿意多包容几分，不是勉强的妥协，而是打心底觉得"对她好就是对的"。她也不会理所当然——土厚载万物，她接得住他的好，也兜得住他的所有时刻。'],
])

# 日主关系图
t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
c0, c1 = t.rows[0].cells[0], t.rows[0].cells[1]
bgx(c0, BG); bgx(c1, BG)
for line in ['男\n丁火 ──生──→ 戊土', '烛火          城墙', '温而不灼      稳而不移', '细腻敏感      大气包容']:
    pa = c0.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=10, c=BK)
for line in ['', '他点亮她的世界    她安定他的漂泊', '她给他归属感      他给她存在感', '']:
    pa = c1.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=10, c=BK)
doc.add_paragraph()

# 维度三
sub_sec('🟡 维度三：年柱 · 同纳音 —— ⭐⭐⭐⭐ 深厚共振')
info_card([
    ['关系', '男丙子 涧下水 · 女丁丑 涧下水 — 完全相同'],
    ['解读', '年柱纳音代表家世根基。同属涧下水，意味着你们来自同一座山、同一条溪——价值观底层是相通的。很多事情不需要解释就能达成一致。'],
    ['现实', '双方原生家庭相处模式天然接近，长辈之间不会有根本性冲突。'],
])

t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
c = t.rows[0].cells[0]; bgx(c, BG)
for line in ['男方 丙子（溪涧流水，灵活流动）', '           ↘', '           同出一源 → 涧下水 ← 根基相通', '           ↗', '女方 丁丑（同源之水，更深更厚）']:
    pa = c.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=10, c=BK)
doc.add_paragraph()

# 维度四
sub_sec('🟡 维度四：月柱 · 申寅六冲 —— ⭐⭐ 需要管理')
info_card([
    ['关系', '男月支 申(猴) ——相冲—— 女月支 寅(虎)'],
    ['解读', '月柱代表父母宫和处事方式。申寅冲是整个合盘中唯一需要认真对待的张力，会在涉及双方家庭、事业决策、生活节奏时产生不同意见。'],
    ['正面效应', '申寅冲是"驿马冲"——不是破坏性的，而是动力型的。你们的感情不会沉闷，生活一直有新鲜感和闯劲。两个人都不是安于现状的人。'],
    ['管理方案', '① 大事提前沟通，别让两家直接对线\n② 申月(7-8月)和寅月(1-2月)情绪偏急，注意避雷\n③ 她的壬水月干润你的丙火——你定调、她落地，正好互补'],
])

# 维度五
sub_sec('🟡 维度五：配偶宫 · 酉戌相害 —— ⭐⭐⭐ 节奏弹性')
info_card([
    ['关系', '男日支 酉(鸡) ——相害—— 女日支 戌(狗)'],
    ['解读', '酉戌相害是六害中最轻的一种。不是冲、不是刑、不是穿——只是你们的生活节奏天然保留了一块弹性空间。一个偏晨型、一个偏夜型，让这个家始终有人在清醒地守护着。'],
    ['现实', '他可能有点讲究、注意细节（酉金）；她偏向随性、抓大放小（戌土）。小事上会有碎碎念——但讨论完就翻篇，从不记仇。'],
    ['程度', '低。酉戌害不会伤及感情根基，是生活中自然的小调剂。'],
])

page_break()

# ═══════════════════════════════════════════════════════
#  四、合盘得分卡
# ═══════════════════════════════════════════════════════

sec('四', '合盘得分卡')

t = doc.add_table(rows=8, cols=5); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['维度', '权重', '评分', '加权', '评价']
for i, h in enumerate(hdrs):
    c = t.rows[0].cells[i]; bgx(c, GOLD); tcell(c, h, sz=10, b=True, c=WH)

rows_data = [
    ['生肖六合', '★★★★★', '1.00', '0.25', '满分'],
    ['日主相生', '★★★★★', '1.00', '0.25', '完美'],
    ['年柱同源', '★★★★☆', '0.80', '0.20', '深厚'],
    ['月柱相冲', '★★☆☆☆', '0.40', '0.08', '可管理'],
    ['配偶宫害', '★★★☆☆', '0.60', '0.12', '轻微'],
]
for ri, row in enumerate(rows_data):
    for ci, v in enumerate(row):
        c = t.rows[ri+1].cells[ci]
        if ri % 2 == 0: bgx(c, BG)
        tcell(c, v, sz=9.5)

# 合并综合行
c_sum = t.rows[6].cells[0]
for i in range(1,5): c_sum.merge(t.rows[6].cells[i])
bgx(c_sum, BG)
tcell(c_sum, '综合  0.90 / 1.00  →  5.00 / 5.00  🟡🟡🟡🟡🟡 | 上等婚配', sz=11, b=True, c=DARK)

c_foot = t.rows[7].cells[0]
for i in range(1,5): c_foot.merge(t.rows[7].cells[i])
tcell(c_foot, '加权：生肖 25% + 日主 25% + 年柱 20% + 月柱 20% + 配偶宫 10%', sz=8.5, c=GR)

t.columns[0].width = Cm(3); t.columns[1].width = Cm(3); t.columns[2].width = Cm(2)
t.columns[3].width = Cm(2); t.columns[4].width = Cm(3)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  五、五行互补
# ═══════════════════════════════════════════════════════

sec('五', '五行互补图谱')

new_table(
    ['元素', '男方', '女方', '互补关系'],
    [
        ['金', '████ 20%', '██████ 28%', '女方金足，降男方火'],
        ['木', '██ 10%', '██ 10%', '男方甲木正印，能疏女方土'],
        ['水', '████ 20%', '███ 15%', '涧水下同源，根基相通'],
        ['火', '████████ 40% (偏旺)', '███ 12%', '男方火旺需降，女方恰好不燥'],
        ['土', '██ 10%', '████████ 35% (偏厚)', '女方土厚需松，男方恰好不硬'],
    ],
    col_w=[2, 4, 4, 6]
)

acres('她是他命里的"降温系统"    他是她命里的"松土铲"', sz=10.5, b=True, c=DARK)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  六、太极图
# ═══════════════════════════════════════════════════════

sec('六', '你们的太极图')

t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER

# 左 — 男方
cl = t.rows[0].cells[0]; bgx(cl, 'FFF5EB')
tcell(cl, '丁火（男方）', sz=11.5, b=True, c=DARK)
for line in ['', '烛火温而不灼', '情细而深', '甲木远见', '', '想得多，思虑周全', '他的火需要她来降温', '她是他命里的安定感']:
    pa = cl.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=9.5, c=BK if line else GR)

# 中 — 互补
cc = t.rows[0].cells[1]; bgx(cc, BG)
tcell(cc, '互 补', sz=12, b=True, c=DARK)
for line in ['', '火 生 土', '', '她急时 → 他退一步', '他倦时 → 她撑一把', '讨论时当面清', '日常零磨合']:
    pa = cc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=10, b=(line=='火 生 土'), c=DARK if line else GR)

# 右 — 女方
cr = t.rows[0].cells[2]; bgx(cr, 'FFF5EB')
tcell(cr, '戊土（女方）', sz=11.5, b=True, c=DARK)
for line in ['', '城墙稳而不移', '气大而厚', '酉金精细', '', '认定了的事全力以赴', '她的土需要他来松解', '他是她命里的小太阳']:
    pa = cr.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pa.add_run(line); fr(r, sz=9.5, c=BK if line else GR)

t.columns[0].width = Cm(5.5); t.columns[1].width = Cm(5); t.columns[2].width = Cm(5.5)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  七、婚配细节建议
# ═══════════════════════════════════════════════════════

sec('七', '婚配细节建议')

new_table(
    ['维度', '说明'],
    [
        ['沟通共振', '申寅月（公历 1-2 月 / 7-8 月）情绪偏急，大事别在这两个时段谈'],
        ['财务分工', '你的丁火配甲木正印——擅长长线布局，投资规划、资产配置由你定方向。她的戊土配酉金伤官——精于落地执行，日常收支、生活开销归她把细节。一个管战略、一个管战壕。'],
        ['讨论模式', '申寅冲 → 能量来得快去得快。不要冷战，当面说开就翻篇。火生土的关系决定了——谁先伸出手，另一方就会握住。'],
        ['养育风格', '你甲木正印在时柱 → 宠孩子冠军。她戊土+庚申食神 → 立规矩担当。慈父+严母，完美分工。'],
        ['未来展望', '不存在"痒"的基础。酉戌害的节奏弹性，恰好让日常始终保持新鲜感——不攒情绪，每一天都是新的。'],
        ['最佳化解', '多去水边一起走走。你们的年柱涧下水同源——水是你们的共同源头，能化解所有地支冲害。'],
    ],
    col_w=[2.5, 13.5]
)

page_break()

# ═══════════════════════════════════════════════════════
#  八、判词
# ═══════════════════════════════════════════════════════

sec('八', '判词')
gold_box([
    '丙子丁丑，涧下同源',
    '丁火戊土，生生不息',
    '子丑六合，天作之合',
    '寅申虽冲，化动力也',
    '酉戌弹性，新鲜常在',
    '',
    '火暖厚土千年固',
    '水归同源万里长',
    '',
    '上等婚配，宜家宜室',
], sz=12)

# ═══════════════════════════════════════════════════════
#  九、每月吉日速查
# ═══════════════════════════════════════════════════════

sec('九', '每月吉日速查 · 避冲指南')

new_table(
    ['月份', '日支避', '原因', '特别提醒'],
    [
        ['1月', '寅日、午日', '寅月+寅申冲，午冲子', '双重敏感月，慎选'],
        ['2月', '卯日、午日', '卯刑子', '春节叠加，等正月十五后'],
        ['3月', '午日', '冲子', '回暖，择日窗口多'],
        ['4月', '午日、未日', '冲子、害子', ''],
        ['5月', '午日', '冲子', ''],
        ['6月', '午日', '冲子', ''],
        ['7月', '寅日、午日', '申月+寅申冲', '情绪敏感月，避开寅日'],
        ['8月', '酉日、午日', '酉酉自刑，午冲子', '申月持续，慎选'],
        ['9月', '午日', '冲子', '窗口开阔'],
        ['10月', '午日', '冲子', '秋高气爽，最佳月份'],
        ['11月', '午日', '冲子', '同样优质'],
        ['12月', '午日', '冲子', '天气偏冷'],
    ],
    col_w=[1.5, 2.5, 4, 8]
)

ap('每月避开午日（冲男鼠）是底线。寅日、未日、酉日按需避开。', sz=9, c=GR)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════
#  十、报告信息
# ═══════════════════════════════════════════════════════

sec('十', '报告信息')
info_card([
    ['分析日期', '2026年6月8日'],
    ['数据来源', '男方提供（1996.8.27 辰时）· 女方提供（1997.2.24 申时）'],
    ['分析方法', '八字四柱合盘法（生肖/日主/年柱/月柱/配偶宫 五维）'],
    ['参考体系', '子平八字 + 民间合婚传统 + 五行生克'],
])

doc.add_paragraph(); doc.add_paragraph()

# ─── 结尾 ───
acres('八字描摹的是缘分的底色；婚姻的画面，终究由两个人握着同一支笔画出。', sz=10, c=GR)
ap()
acres('🐭 ❤️ 🐮', sz=16, b=True, c=GOLD)
ap()
acres('本报告供个人参考，不构成任何形式的建议或承诺。', sz=8, c='AAAAAA')

# ─── 保存 ───
try:
    os.remove(OUTPUT)
except:
    pass
doc.save(OUTPUT)
print(f'OK: {OUTPUT}')
